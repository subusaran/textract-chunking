import docx
import os

def create_dummy_docx(filename):
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    
    doc.add_paragraph('This is a regular paragraph with some text.')
    
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Another paragraph in section 1.')
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header 1'
    hdr_cells[1].text = 'Header 2'
    hdr_cells[2].text = 'Header 3'
    
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'R1C1'
    row1_cells[1].text = 'R1C2'
    row1_cells[2].text = 'R1C3'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'R2C1'
    row2_cells[1].text = 'R2C2'
    row2_cells[2].text = 'R2C3'
    
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    create_dummy_docx("test_doc.docx")
